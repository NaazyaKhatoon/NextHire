import io
from typing import Optional
import pypdf
import docx

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return ""

def extract_text_from_docx(docx_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        return "\n".join(full_text)
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return ""
